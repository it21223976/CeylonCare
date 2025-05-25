#!/usr/bin/env python3
"""
Train the SriLankanFoodMLSystem end-to-end
Usage:
    python train_ml.py <categorization_docx> <recipes_docx>
"""
import os
import sys
import json
import re

import pandas as pd
import numpy as np
import tensorflow as tf
from tensorflow.keras.models import Model
from tensorflow.keras.layers import Dense, Dropout, Input
from sklearn.preprocessing import StandardScaler, LabelEncoder
from sklearn.cluster import KMeans
from sklearn.ensemble import RandomForestClassifier
import joblib
from docx import Document

# Suppress TensorFlow warnings
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '2'

class SriLankanFoodMLSystem:
    def __init__(self):
        self.models_dir = 'ml_models'
        self.data_dir   = 'ml_data'
        self._make_dirs()

    def _make_dirs(self):
        os.makedirs(self.models_dir, exist_ok=True)
        os.makedirs(self.data_dir, exist_ok=True)

    def extract_data_from_docx(self, cat_file, rec_file):
        print(f"\n📂 Loading categorization file: {cat_file}")
        doc_cat = Document(cat_file)
        print("  → Total paragraphs in categorization doc:", len(doc_cat.paragraphs))
        cats = self._extract_categories(cat_file)
        print("🔍 Categories found:", {k: len(v) for k, v in cats.items()})

        print(f"\n📂 Loading recipes file: {rec_file}")
        doc_rec = Document(rec_file)
        print("  → Total paragraphs in recipes doc:", len(doc_rec.paragraphs))
        recs = self._extract_recipes(rec_file)
        print("🔍 Recipes found:", len(recs), "– sample keys:", list(recs.keys())[:5])

        df = self._build_dataset(cats, recs)
        print("🔍 Built dataset rows:", len(df))
        out_csv = os.path.join(self.data_dir, 'sri_lankan_foods.csv')
        df.to_csv(out_csv, index=False)
        print(f"✅ Saved dataset to {out_csv}")
        return df

    def _extract_categories(self, path):
            doc = Document(path)
            cats = {'high_blood_pressure': [], 'diabetes': [], 'both': []}
            cur = None

            for p in doc.paragraphs:
                t = p.text.strip()
                if not t:
                    continue

                low = t.lower()
                if low == 'diabetics':
                    cur = 'diabetes'; continue
                if low == 'high blood pressure':
                    cur = 'high_blood_pressure'; continue
                if low == 'both':
                    cur = 'both'; continue

                # once we're within a section, treat any non‐header line as an entry
                if cur:
                    # strip leading numbering (e.g. "1. Foo") and trailing parens
                    name = re.sub(r'^\d+\.\s*', '', t)
                    name = re.sub(r'\s*\([^)]*\)\s*$', '', name).strip()
                    cats[cur].append(name)

            return cats    


    def _extract_recipes(self, path):
        """Extract recipes from DOCX file"""
        paras = [p.text.strip() for p in Document(path).paragraphs if p.text.strip()]
        print("\n   [DEBUG] _extract_recipes got", len(paras), "non-empty paragraphs")
        recipes = {}
        current = None
        section = None

        for idx, line in enumerate(paras):
            print(f"     L{idx:02d}:", repr(line))
            
            # Match recipe headers (numbered format)
            m = re.match(r'^(\d+)\.\s*(.+)', line)
            if m:
                if current:
                    recipes[current['name']] = current
                    print(f"       → Saved recipe: {current['name']}")
                
                recipe_name = m.group(2).strip()
                # Clean parenthetical descriptions
                recipe_name = re.sub(r'\s*\([^)]*\)\s*$', '', recipe_name)
                
                current = {
                    'name': recipe_name,
                    'ingredients': [],
                    'instructions': [],
                    'portion_size': '1 serving'  # Default portion size
                }
                print("       → New recipe start:", current['name'])
                section = None
                continue

            # Section headers
            if line.lower().startswith('ingredients'):
                section = 'ingredients'
                print("       → Entering INGREDIENTS section")
                continue
            if line.lower().startswith('preparation') or line.lower().startswith('instructions'):
                section = 'preparation'
                print("       → Entering PREPARATION section")
                continue
            if line.lower().startswith('portion size'):
                pm = re.search(r'Portion Size[:\s]*(.+)', line, re.IGNORECASE)
                if pm:
                    current['portion_size'] = pm.group(1).strip()
                    print("       → Portion size:", current['portion_size'])
                section = None
                continue

            # Ingredients (bullet points)
            if section == 'ingredients' and (line.startswith('•') or line.startswith('-')):
                ingr = line.lstrip('•-\t ').strip()
                if current:
                    current['ingredients'].append(ingr)
                    print("         • Ingredient:", ingr)
                continue

            # Instructions (numbered)
            if section == 'preparation' and re.match(r'^\d+\.', line):
                instr = re.sub(r'^\d+\.\s*', '', line).strip()
                if current:
                    current['instructions'].append(instr)
                    print("         • Instruction:", instr)
                continue

        # Save last recipe
        if current:
            recipes[current['name']] = current
            print(f"       → Saved last recipe: {current['name']}")
        
        print("   [DEBUG] Recipe keys:", list(recipes.keys()))
        return recipes

    def _build_dataset(self, cats, recs):
        """Build dataset from categories and recipes"""
        nut_db = self._nutrition_db()
        rows = []
        fid = 1
        
        for cat, names in cats.items():
            for nm in names:
                # Try to find recipe, otherwise create default
                rc = recs.get(nm)
                if not rc:
                    # Try alternative names or partial matches
                    matched = False
                    for rec_name in recs.keys():
                        if nm.lower() in rec_name.lower() or rec_name.lower() in nm.lower():
                            rc = recs[rec_name]
                            matched = True
                            print(f"   [MATCH] Found recipe '{rec_name}' for food '{nm}'")
                            break
                    
                    if not matched:
                        rc = {
                            'name': nm,
                            'ingredients': [f"Traditional {nm} method"],
                            'instructions': ["Follow traditional method"],
                            'portion_size': "1 serving"
                        }
                        print(f"   [DEFAULT] Using default recipe for '{nm}'")
                
                nut = self._estimate_nutrition(rc, nm, nut_db)
                row = {
                    'id': fid,
                    'name': nm,
                    'health_category': cat,
                    'recipe_category': '',  # Empty for now
                    'ingredients': json.dumps(rc['ingredients']),
                    'instructions': json.dumps(rc['instructions']),
                    'portion_size': rc['portion_size'],
                    **nut
                }
                rows.append(row)
                fid += 1
                
        return pd.DataFrame(rows)

    def _nutrition_db(self):
        """Expanded nutrition database for Sri Lankan foods"""
        return {
            # Grains and legumes
            'kurakkan': {'protein': 7.3, 'fat': 1.3, 'carbs': 72.6, 'potassium': 408, 'fiber': 11.5, 'sodium': 0.1, 'magnesium': 2.8},
            'kawpi': {'protein': 8.7, 'fat': 0.5, 'carbs': 20.9, 'potassium': 278, 'fiber': 6.3, 'sodium': 0.2, 'magnesium': 1.8},
            'mung': {'protein': 7.0, 'fat': 0.4, 'carbs': 19.1, 'potassium': 266, 'fiber': 4.1, 'sodium': 0.1, 'magnesium': 1.5},
            'chickpea': {'protein': 8.9, 'fat': 2.6, 'carbs': 27.4, 'potassium': 291, 'fiber': 7.6, 'sodium': 0.2, 'magnesium': 2.4},
            'brown rice': {'protein': 2.6, 'fat': 0.9, 'carbs': 23.0, 'potassium': 43, 'fiber': 1.8, 'sodium': 0.1, 'magnesium': 0.8},
            
            # Vegetables
            'coconut': {'protein': 3.3, 'fat': 33.5, 'carbs': 15.2, 'potassium': 356, 'fiber': 9.0, 'sodium': 0.2, 'magnesium': 3.2},
            'pumpkin': {'protein': 1.0, 'fat': 0.1, 'carbs': 6.5, 'potassium': 340, 'fiber': 0.5, 'sodium': 0.1, 'magnesium': 0.8},
            'eggplant': {'protein': 1.0, 'fat': 0.2, 'carbs': 5.9, 'potassium': 229, 'fiber': 3.0, 'sodium': 0.1, 'magnesium': 0.7},
            'tomato': {'protein': 0.9, 'fat': 0.2, 'carbs': 3.9, 'potassium': 237, 'fiber': 1.2, 'sodium': 0.1, 'magnesium': 0.6},
            'cabbage': {'protein': 1.3, 'fat': 0.1, 'carbs': 5.8, 'potassium': 170, 'fiber': 2.5, 'sodium': 0.2, 'magnesium': 0.7},
            'yam': {'protein': 1.5, 'fat': 0.1, 'carbs': 27.9, 'potassium': 816, 'fiber': 4.1, 'sodium': 0.1, 'magnesium': 1.2},
            'manioc': {'protein': 1.4, 'fat': 0.3, 'carbs': 38.1, 'potassium': 271, 'fiber': 1.8, 'sodium': 0.1, 'magnesium': 1.1},
            'breadfruit': {'protein': 1.1, 'fat': 0.2, 'carbs': 27.1, 'potassium': 490, 'fiber': 4.9, 'sodium': 0.1, 'magnesium': 1.5},
            
            # Leafy greens
            'gotu kola': {'protein': 1.6, 'fat': 0.2, 'carbs': 4.6, 'potassium': 345, 'fiber': 2.1, 'sodium': 0.1, 'magnesium': 1.9},
            'murunga': {'protein': 2.1, 'fat': 0.1, 'carbs': 3.7, 'potassium': 259, 'fiber': 2.0, 'sodium': 0.1, 'magnesium': 1.7},
            'mugunuwanna': {'protein': 2.9, 'fat': 0.4, 'carbs': 5.2, 'potassium': 558, 'fiber': 2.2, 'sodium': 0.2, 'magnesium': 2.1},
            'thebu': {'protein': 2.5, 'fat': 0.3, 'carbs': 4.8, 'potassium': 420, 'fiber': 2.5, 'sodium': 0.1, 'magnesium': 2.0},
            'aguna': {'protein': 2.2, 'fat': 0.2, 'carbs': 4.5, 'potassium': 380, 'fiber': 2.3, 'sodium': 0.1, 'magnesium': 1.8},
            
            # Fruits
            'papaya': {'protein': 0.5, 'fat': 0.3, 'carbs': 10.8, 'potassium': 182, 'fiber': 1.7, 'sodium': 0.1, 'magnesium': 0.9},
            'banana': {'protein': 1.1, 'fat': 0.3, 'carbs': 22.8, 'potassium': 358, 'fiber': 2.6, 'sodium': 0.1, 'magnesium': 1.3},
            'nelli': {'protein': 0.9, 'fat': 0.6, 'carbs': 10.2, 'potassium': 198, 'fiber': 4.3, 'sodium': 0.1, 'magnesium': 1.0},
            'starfruit': {'protein': 1.0, 'fat': 0.3, 'carbs': 6.7, 'potassium': 133, 'fiber': 2.8, 'sodium': 0.1, 'magnesium': 0.8},
            
            # Fish and others
            'fish': {'protein': 20.0, 'fat': 1.0, 'carbs': 0.0, 'potassium': 350, 'fiber': 0.0, 'sodium': 0.5, 'magnesium': 2.5},
            'sprats': {'protein': 19.0, 'fat': 2.5, 'carbs': 0.0, 'potassium': 320, 'fiber': 0.0, 'sodium': 0.8, 'magnesium': 2.2},
        }

    def _estimate_nutrition(self, recipe, name, db):
        """Estimate nutritional values based on food name and ingredients"""
        # Base nutritional values
        base = {
            'protein': 2.0,
            'fat': 1.0,
            'carbs': 15.0,
            'potassium': 300.0,
            'fiber': 2.5,
            'sodium': 0.5,
            'magnesium': 0.8
        }
        
        # Combine recipe ingredients and name for matching
        txt = ' '.join(recipe['ingredients']).lower() + ' ' + name.lower()
        
        # Match ingredients from nutrition database
        matched_nutrients = []
        for ingr, vals in db.items():
            if ingr in txt:
                matched_nutrients.append((ingr, vals))
                print(f"      [NUTRITION] Matched '{ingr}' in '{name}'")
        
        # If we have matches, use weighted average
        if matched_nutrients:
            # Reset base to zero for averaging
            for k in base:
                base[k] = 0
            
            # Average the matched nutrients
            for ingr, vals in matched_nutrients:
                weight = 1.0 / len(matched_nutrients)
                for k in base:
                    base[k] += vals.get(k, 0) * weight
        
        # Calculate derived values
        base['calories'] = int(base['protein']*4 + base['fat']*9 + base['carbs']*4)
        base['cooking_time'] = 20  # Default cooking time
        
        # Calculate health score based on nutritional profile
        health_score = 60  # Base score
        if base['fiber'] > 4:
            health_score += 15
        if base['potassium'] > 400:
            health_score += 15
        if base['sodium'] < 0.5:
            health_score += 10
        base['health_score'] = min(100, health_score)
        
        # Estimate glycemic index based on food type
        if 'brown rice' in txt or 'kurakkan' in txt:
            base['glycemic_index'] = 55
        elif 'white' in txt and 'rice' in txt:
            base['glycemic_index'] = 70
        elif any(fruit in txt for fruit in ['papaya', 'banana', 'nelli']):
            base['glycemic_index'] = 45
        elif any(veg in txt for veg in ['cabbage', 'eggplant', 'tomato']):
            base['glycemic_index'] = 35
        else:
            base['glycemic_index'] = 50  # Default
        
        return base

    def train_all_models(self, df):
        """Train all ML models"""
        print("\n🤖 Training models…")
        
        # Feature columns
        feats = [
            'protein', 'fat', 'carbs', 'potassium', 'fiber', 'sodium', 'magnesium',
            'calories', 'cooking_time', 'glycemic_index'
        ]
        
        if df.empty:
            raise ValueError("Dataset is empty—check your DOCX parsing.")
        
        X = df[feats].values
        print(f"   [INFO] Training with {len(X)} samples and {len(feats)} features")

        # 1. Content-based model
        print("\n   📊 Training content-based model...")
        scaler_cb = StandardScaler().fit(X)
        X_cb = scaler_cb.transform(X)
        
        # Create labels
        y_bp = (df['health_category'].isin(['high_blood_pressure', 'both'])).astype(int)
        y_db = (df['health_category'].isin(['diabetes', 'both'])).astype(int)
        y_hs = df['health_score'] / 100.0
        
        # Build neural network
        inp = Input(shape=(X.shape[1],))
        h1 = Dense(64, activation='relu')(inp)
        d1 = Dropout(0.3)(h1)
        h2 = Dense(32, activation='relu')(d1)
        d2 = Dropout(0.2)(h2)
        
        bp_out = Dense(1, activation='sigmoid', name='blood_pressure')(d2)
        db_out = Dense(1, activation='sigmoid', name='diabetes')(d2)
        hs_out = Dense(1, activation='sigmoid', name='health_score')(d2)

        model_cb = Model(inp, [bp_out, db_out, hs_out])
        model_cb.compile(
            optimizer='adam',
            loss={
                'blood_pressure': 'binary_crossentropy',
                'diabetes': 'binary_crossentropy',
                'health_score': 'mse'
            },
            metrics={
                'blood_pressure': ['accuracy'],
                'diabetes': ['accuracy'],
                'health_score': ['mae']
            }
        )

        # Train the model
        history = model_cb.fit(
            X_cb, [y_bp, y_db, y_hs],
            epochs=50,  # Increased epochs
            batch_size=8,
            validation_split=0.2,
            verbose=1
        )
        
        # Save model and scaler
        model_cb.save(os.path.join(self.models_dir, 'content_based_model.h5'))
        joblib.dump(scaler_cb, os.path.join(self.models_dir, 'content_scaler.pkl'))
        print("   ✅ Content-based model trained")

        # 2. Clustering model
        print("\n   📊 Training clustering model...")
        scaler_cl = StandardScaler().fit(X)
        X_cl = scaler_cl.transform(X)
        
        # Determine optimal number of clusters
        k = min(8, max(3, len(df) // 5))
        print(f"   [INFO] Using {k} clusters")
        
        kmeans = KMeans(n_clusters=k, random_state=42, n_init=10).fit(X_cl)
        joblib.dump(kmeans, os.path.join(self.models_dir, 'kmeans.pkl'), protocol=4)
        joblib.dump(scaler_cl, os.path.join(self.models_dir, 'kmeans_scaler.pkl'))
        print("   ✅ Clustering model trained")

        # 3. Health classifier
        print("\n   📊 Training health classifier...")
        le = LabelEncoder().fit(df['health_category'])
        y = le.transform(df['health_category'])
        
        rf = RandomForestClassifier(
            n_estimators=100,  # Increased trees
            max_depth=10,
            random_state=42
        ).fit(X, y)
        
        joblib.dump(rf, os.path.join(self.models_dir, 'rf_health.pkl'), protocol=4)
        joblib.dump(le, os.path.join(self.models_dir, 'health_le.pkl'))
        print("   ✅ Health classifier trained")

        # 4. Save metadata
        print("\n   📊 Saving metadata...")
        meta = {
            'total_rows': len(df),
            'categories': df['health_category'].value_counts().to_dict(),
            'features': feats,
            'foods_by_category': {
                cat: df[df['health_category'] == cat]['name'].tolist()
                for cat in df['health_category'].unique()
            },
            'training_date': pd.Timestamp.now().isoformat()
        }
        
        with open(os.path.join(self.models_dir, 'metadata.json'), 'w') as f:
            json.dump(meta, f, indent=2)
        
        print("\n✅ All models trained successfully!")
        
        # Print summary
        print("\n📊 Training Summary:")
        print(f"   • Total foods: {len(df)}")
        for cat, count in df['health_category'].value_counts().items():
            print(f"   • {cat}: {count} foods")
        
        # Return success result
        result = {
            "success": True,
            "models_trained": {
                "content_based": True,
                "clustering": True,
                "health_classifier": True
            },
            "total_foods": len(df),
            "categories": df['health_category'].value_counts().to_dict()
        }
        
        # Print JSON result for the Node.js backend
        print("\n" + json.dumps(result))
        sys.stdout.flush()

def main():
    if len(sys.argv) != 3:
        print("Usage: python train_ml.py <categorization.docx> <recipes.docx>")
        sys.exit(1)

    cat_file = sys.argv[1]
    rec_file = sys.argv[2]

    try:
        ml_system = SriLankanFoodMLSystem()
        df = ml_system.extract_data_from_docx(cat_file, rec_file)
        ml_system.train_all_models(df)
    except Exception as e:
        error_result = {
            "success": False,
            "error": str(e)
        }
        print(json.dumps(error_result))
        sys.exit(1)

if __name__ == "__main__":
    main()